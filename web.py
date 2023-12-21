import streamlit as st
import csv
from openpyxl import Workbook
from docx import Document
from docx.shared import Cm
import base64
from io import BytesIO
import os

def csv_to_list(file_path):
    data_list = []
    with open(file_path, 'r', encoding='utf-8') as csv_file:
        csv_reader = csv.reader(csv_file)
        for row in csv_reader:
            data_list.append(row)
    return data_list

def append_to_excel(file_io, data):
    wb = Workbook()
    ws = wb.active
    for row in data:
        ws.append(row)
    wb.save(file_io)

def append_to_word(doc, data, filename):
    doc.add_heading(f'Members from {filename}', level=1)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    table = doc.add_table(rows=1, cols=5)  # Table with 5 columns
    table.style = 'Table Grid'

    # Set column headers
    hdr_cells = table.rows[0].cells
    columns = ['Sr. No', 'Name', 'Address', 'Area', 'Ms. No']  # Adjusted column names
    for col_index, cell in enumerate(columns):
        hdr_cells[col_index].text = str(cell)
        hdr_cells[col_index].paragraphs[0].alignment = 1
        hdr_cells[col_index].paragraphs[0].runs[0].bold = True

    for row in data[0:]:
        row_cells = table.add_row().cells
        for col_index, cell in enumerate(row[:5]):  # Only first five columns
            row_cells[col_index].text = str(cell)

    # Adjust width of columns
    column_widths = [Cm(1.5), None, Cm(5), Cm(4), Cm(1.5)]  # Column widths in cm
    table.autofit = True
    for col_index, width in enumerate(column_widths):
        if width:
            table.columns[col_index].width = width

    # Calculate width for Column 2 with remaining space
    remaining_width = sum((col.width for col in table.columns if col.width is not None))
    col_2_width = table.columns[1].width = remaining_width

def filter_and_append(csv_file_path, words_list):
    data_to_append = []
    with open(csv_file_path, 'r', encoding='utf-8') as csv_file:
        csv_reader = csv.reader(csv_file)
        for row in csv_reader:
            row_lower = [col.lower().replace(" ", "") for col in row]  # Removing spaces from CSV data
            for word in words_list:
                if any(word.lower().replace(" ", "") in cell for cell in row_lower[2:4]):
                    data_to_append.append(row)
                    break  # Break if any word is found in the row to avoid duplicates
    return data_to_append

def main():
    st.title('CSV Data Processing')

    csv_file_path = 'new2.csv'

    # Check if the CSV file exists
    if not os.path.exists(csv_file_path):
        st.error("CSV file not found at the specified path.")
        st.stop()

    user_input = st.text_input('Enter words separated by space (Type "exit" to quit):')
    if user_input.lower() == 'exit':
        st.stop()

    words_list = user_input.split()

    file_name_input = st.text_input('Enter a file name for Excel and Word files:')
    if st.button('Submit'):
        folder_name1 = 'print/excel'
        folder_name2 = 'print/word'
        excel_file_path = f'{folder_name1}/{file_name_input}.xlsx'
        word_file_path = f'{folder_name2}/{file_name_input}.docx'

        filtered_data = filter_and_append(csv_file_path, words_list)
        if filtered_data:
            try:
                # Prepare Excel file in memory
                excel_io = BytesIO()
                append_to_excel(excel_io, filtered_data)
                excel_io.seek(0)

                # Prepare Word file in memory
                word_io = BytesIO()
                doc = Document()
                append_to_word(doc, filtered_data, file_name_input)
                doc.save(word_io)
                word_io.seek(0)

                st.write(f"Data for '{', '.join(words_list)}' is ready to download.")

                # Download Excel file
                excel_bytes = excel_io.read()
                b64_excel = base64.b64encode(excel_bytes).decode()
                st.markdown(f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64_excel}" download="{file_name_input}.xlsx">Download Excel File</a>', unsafe_allow_html=True)

                # Download Word file
                word_bytes = word_io.read()
                b64_word = base64.b64encode(word_bytes).decode()
                st.markdown(f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_word}" download="{file_name_input}.docx">Download Word File</a>', unsafe_allow_html=True)
                
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
        else:
            st.write(f"No matching data found for '{', '.join(words_list)}'.")

if __name__ == '__main__':
    main()
